import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pandas as pd
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter

class MultimodalParser:
    def __init__(self, tesseract_cmd_path=None, threshold=2000):
        # ตั้งค่า path สำหรับ Tesseract (ถ้าจำเป็น)
        if tesseract_cmd_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd_path
        
        self.threshold = threshold
        # ใช้ Langchain text splitter สำหรับการหั่นข้อความ (Chunking)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def _get_document_type(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ['.docx']:
            return "DOCX"
        elif ext in ['.pdf']:
            return "PDF"
        elif ext in ['.png', '.jpg', '.jpeg']:
            return "IMG"
        elif ext in ['.xls', '.xlsx']:
            return "Excel"
        else:
            return "UNKNOWN"

    def parse_document(self, file_path: str) -> list:
        T = [] # Initialize T
        doc_type = self._get_document_type(file_path)

        # ---------------------------------------------------------
        # 1. Parse DOCX
        # ---------------------------------------------------------
        if doc_type == "DOCX":
            doc = docx.Document(file_path)
            
            # Extract Text and Tables
            # (Note: python-docx ไม่รองรับการดึงรูปภาพแบบ inline ได้ง่ายๆ 
            # ในระดับโปรดักชันมักจะใช้ไลบรารีอื่นช่วยดึงรูปจาก docx แยกต่างหาก)
            for element in doc.elements: # This is conceptual, python-docx uses paragraphs/tables separately
                pass 
            
            # ดึงข้อความจาก Paragraph
            for para in doc.paragraphs:
                if para.text.strip():
                    T.append(para.text.strip())
            
            # ดึงตารางและแปลงเป็น Markdown
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    markdown_table = df.to_markdown(index=False)
                    T.append(markdown_table)

        # ---------------------------------------------------------
        # 2. Parse PDF
        # ---------------------------------------------------------
        elif doc_type == "PDF":
            pdf_document = fitz.open(file_path)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Extract Text
                page_text = page.get_text()
                if page_text.strip():
                    T.append(page_text.strip())
                
                # Extract Images
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # เซฟรูปชั่วคราวเพื่อทำ OCR
                    temp_img_path = f"temp_image_{page_num}_{img_index}.png"
                    with open(temp_img_path, "wb") as f:
                        f.write(image_bytes)
                    
                    # OCR
                    img_text = pytesseract.image_to_string(Image.open(temp_img_path), lang='tha+eng')
                    if img_text.strip():
                        T.append(img_text.strip())
                    
                    os.remove(temp_img_path) # ลบไฟล์ชั่วคราว

        # ---------------------------------------------------------
        # 3. Parse Image (IMG)
        # ---------------------------------------------------------
        elif doc_type == "IMG":
            img_text = pytesseract.image_to_string(Image.open(file_path), lang='tha+eng')
            if img_text.strip():
                T.append(img_text.strip())

        # ---------------------------------------------------------
        # 4. Parse Excel
        # ---------------------------------------------------------
        elif doc_type == "Excel":
            # อ่านทุก Sheet ใน Excel
            excel_data = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, df in excel_data.items():
                # แปลงตารางเป็น Markdown พร้อมชื่อ Sheet
                markdown_table = f"### Sheet: {sheet_name}\n" + df.to_markdown(index=False)
                T.append(markdown_table)

        # ---------------------------------------------------------
        # 5. Split Text By Sections (Chunking)
        # ---------------------------------------------------------
        # รวมความยาวของ text ทั้งหมดใน List เพื่อเช็คเงื่อนไข Threshold
        total_length = sum(len(text) for text in T)
        
        if total_length > self.threshold:
            # รวม Text ทั้งหมดเป็นก้อนเดียวเพื่อเข้ากระบวนการ Split
            joined_text = "\n\n".join(T)
            # ตัดแบ่งเป็น Chunk ย่อยๆ
            T = self.text_splitter.split_text(joined_text)

        return T

# ==========================================
# ตัวอย่างการเรียกใช้งาน
# ==========================================
if __name__ == "__main__":
    parser = MultimodalParser()
    
    # สมมติว่ามีไฟล์เอกสาร
    # result_chunks = parser.parse_document("sample_company_data.pdf")
    
    # for i, chunk in enumerate(result_chunks):
    #     print(f"--- Chunk {i+1} ---")
    #     print(chunk)
    #     print("-" * 20)